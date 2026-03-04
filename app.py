import streamlit as st
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.llms import HuggingFacePipeline
from langchain.chains import RetrievalQA
from langchain.schema import Document

from transformers import pipeline

# -----------------------------
# Text extraction functions
# -----------------------------
def extract_text_from_pdf(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_docx(file):
    doc = DocxDocument(file)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_text_from_txt(file):
    return file.read().decode("utf-8")

# -----------------------------
# Build vectorstore from input
# -----------------------------
def process_input(input_type, input_data):
    documents = []

    # --- URLs ---
    if input_type == "Link":
        for url in input_data:
            if url.strip():
                loader = WebBaseLoader(url)
                docs = loader.load()
                documents.extend(docs)

    # --- Files / Text ---
    elif input_type == "PDF":
        text = extract_text_from_pdf(input_data)
        documents.append(Document(page_content=text))
    elif input_type == "DOCX":
        text = extract_text_from_docx(input_data)
        documents.append(Document(page_content=text))
    elif input_type == "TXT":
        text = extract_text_from_txt(input_data)
        documents.append(Document(page_content=text))
    elif input_type == "Text":
        documents.append(Document(page_content=input_data))

    # Split text into chunks
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)
    chunks = splitter.create_documents([d.page_content for d in documents])

    # Embeddings
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

    # Create FAISS vectorstore
    vectorstore = FAISS.from_documents(chunks, embeddings)
    return vectorstore

# -----------------------------
# RAG QA
# -----------------------------
def answer_question(vectorstore, query, chat_history=None):
    """
    vectorstore : FAISS vectorstore built from documents
    query       : user question
    chat_history: optional conversation history for context
    """

    from transformers import pipeline
    from langchain_community.llms import HuggingFacePipeline
    from langchain.chains import RetrievalQA

    # -----------------------------
    # LLM pipeline
    # -----------------------------
    llm_pipeline = pipeline(
        "text2text-generation",
        model="google/flan-t5-base",  # lightweight model
        max_new_tokens=256,
        temperature=0.3
    )
    llm = HuggingFacePipeline(pipeline=llm_pipeline)

    # -----------------------------
    # RetrievalQA chain
    # -----------------------------
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
        chain_type="stuff",
        return_source_documents=True
    )

    # -----------------------------
    # Query the document
    # -----------------------------
    result = qa_chain(query)

    # -----------------------------
    # Out-of-scope detection
    # -----------------------------
    if not result["source_documents"]:
        return "⚠️ This question is not related to the uploaded document."

    # -----------------------------
    # Lightweight post-processing for "topic" or "summary"
    # -----------------------------
    answer_text = result["result"].strip()

    # Detect if user is asking for topic or summary
    topic_keywords = ["topic", "main idea", "subject", "about", "central theme"]
    summary_keywords = ["summarize", "summary", "brief", "short version"]

    if any(k in query.lower() for k in topic_keywords):
        # Simple heuristic: take first sentence or first 10 words as "topic"
        answer_text = answer_text.split(".")[0]
        if len(answer_text.split()) > 10:
            answer_text = " ".join(answer_text.split()[:10]) + "..."
        answer_text = f"📌 Topic: {answer_text}"

    elif any(k in query.lower() for k in summary_keywords):
        # Summarize: take first 2 sentences (for small model)
        sentences = answer_text.split(".")
        summary = ".".join(sentences[:2]).strip()
        answer_text = f"📝 Summary: {summary}."

    return answer_text
